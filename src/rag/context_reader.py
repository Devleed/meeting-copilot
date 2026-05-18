# =============================================================================
# COMPONENT: context_reader.py
#
# Document loading layer. Reads meeting context files at startup and returns
# their content so the AI assistant knows what the meeting is about.
#
# Class breakdown (one class per responsibility):
#
#   ContextReader        (abstract) — defines the interface contract:
#                          read(filepath) → str. Every concrete reader must
#                          implement this method; callers depend only on this
#                          abstraction, not on pypdf/docx directly.
#
#   TxtContextReader     — reads plain-text (.txt) files.
#   PdfContextReader     — reads PDF files using pypdf.
#   DocxContextReader    — reads Word documents using python-docx.
#
#   ContextReaderFactory — maps file extensions to reader classes.
#                          Adding a new format = one new class + one dict entry;
#                          no existing reader changes (Open/Closed Principle).
#
#   ContextLoader        — public entry point: validates the file exists,
#                          delegates to the factory, and wraps the result in
#                          the [{filename, content}] dict format consumed by
#                          RetrieverBuilder.
#
# SOLID notes:
#   S — each reader handles exactly one file format.
#   O — new formats are supported without modifying existing readers.
#   L — all ContextReader subclasses are fully substitutable for the base class.
#   I — ContextReader exposes only one method (read); no fat interface.
#   D — ContextLoader depends on the ContextReader abstraction, not on pypdf
#       or docx libraries directly.
# =============================================================================

import os   # os — standard library; os.path.exists and os.path.basename / splitext
import sys  # sys — standard library; sys.exit() to terminate on unrecoverable errors

# ABC — Abstract Base Class; provides the machinery to declare abstract classes.
# abstractmethod — decorator that marks a method as required by all subclasses.
# In TS: equivalent to declaring a method in an interface with no body.
from abc import ABC, abstractmethod


class ContextReader(ABC):
    """
    Abstract base class (interface) for all document readers.

    Responsibility: define the contract every concrete reader must fulfil.
    Given a file path, return the document's full text as a plain string.

    Using ABC satisfies the Dependency Inversion Principle: callers
    (ContextLoader, tests) depend on this abstraction rather than on
    specific third-party libraries like pypdf or python-docx.
    """

    @abstractmethod   # @abstractmethod — forces each subclass to implement this method.
                      # Instantiating a subclass that skips it raises TypeError.
    def read(self, filepath: str) -> str:
        """
        Read the file at `filepath` and return its full text content.

        Parameters
        ----------
        filepath : str
            Absolute or relative path to the file on disk.

        Returns
        -------
        str
            The complete text content extracted from the file.
        """
        ...  # ellipsis — conventional placeholder body for abstract methods; never run


class TxtContextReader(ContextReader):
    """
    Reads plain-text (.txt) files.

    Responsibility: open a text file and return its raw string content.
    No format-specific parsing — plain text is returned verbatim.
    """

    def read(self, filepath: str) -> str:
        # `with open(filepath, "r") as f:` — context manager that opens the file
        # in text-read mode and automatically closes it when the block exits,
        # even if an exception is raised inside the block.
        # In TS: equivalent to try { const data = fs.readFileSync(filepath, 'utf8'); }
        # finally { /* file is already closed automatically by the OS */ }
        with open(filepath, "r") as f:
            # f.read() — reads the entire file into a single string in one call.
            # For very large files this loads everything into RAM at once;
            # acceptable here because meeting context documents are typically small.
            return f.read()


class PdfContextReader(ContextReader):
    """
    Reads PDF files (.pdf) using the pypdf library.

    Responsibility: extract text from every page of a PDF and join the
    pages with newlines. Layout, images, and formatting are discarded —
    only raw text content is returned.
    """

    def read(self, filepath: str) -> str:
        # Lazy import — only import pypdf when this reader is actually used.
        # This avoids an ImportError at startup on machines where pypdf is not
        # installed, as long as the user never provides a .pdf file.
        from pypdf import PdfReader  # PdfReader — main class for parsing PDF structure

        # PdfReader(filepath) — opens and parses the binary PDF file.
        # Internally builds a page tree representing the document structure.
        reader = PdfReader(filepath)

        # reader.pages — a list-like sequence of PageObject instances.
        # Generator expression (for ... in ... if ...) — lazy evaluation:
        #   page.extract_text() — extracts the text layer from one page as a string,
        #                         or None if the page is a scanned image with no text layer.
        #   if page.extract_text() — skip None and empty-string pages (scans, blank pages).
        # "\n".join(...) — concatenate all page texts with a newline between pages.
        return "\n".join(
            page.extract_text()          # text from one PDF page
            for page in reader.pages     # iterate over every page in the document
            if page.extract_text()       # skip pages with no extractable text
        )


class DocxContextReader(ContextReader):
    """
    Reads Microsoft Word documents (.docx) using the python-docx library.

    Responsibility: extract paragraph text from a .docx file.
    Tables, images, headers, and footers are not extracted — only body paragraphs.
    """

    def read(self, filepath: str) -> str:
        # Lazy import — only import python-docx when a .docx is actually needed.
        from docx import Document  # Document — main class in the python-docx library

        # Document(filepath) — opens and parses the .docx ZIP/XML structure.
        doc = Document(filepath)

        # doc.paragraphs — a list of Paragraph objects, one per text block.
        # para.text — the plain-text string of one paragraph (no style info).
        # if para.text — skip paragraphs with empty text (blank lines, section breaks).
        # "\n".join(...) — concatenate non-empty paragraphs separated by newlines.
        return "\n".join(
            para.text                # plain text of this paragraph
            for para in doc.paragraphs   # iterate over every paragraph in the document
            if para.text             # skip paragraphs that contain no text
        )


class ContextReaderFactory:
    """
    Factory that maps file extensions to the correct ContextReader subclass.

    Responsibility: decide which reader to instantiate based on a file extension.

    Open/Closed Principle in practice:
      - To support a new format (e.g. .md), write a new ContextReader subclass
        and add one entry to _READERS.
      - No existing reader class needs to change.

    This is the Factory Method pattern: callers ask for a reader by extension
    and receive a concrete instance without knowing which class was constructed.
    """

    # _READERS : class-level dict — extension string → reader class.
    # Using a dict instead of if/elif chains makes adding new formats trivial.
    # Keys are lowercase extension strings including the leading dot.
    _READERS: dict = {
        ".txt":  TxtContextReader,    # plain text files
        ".pdf":  PdfContextReader,    # PDF documents
        ".docx": DocxContextReader,   # Microsoft Word documents
    }

    @classmethod   # @classmethod — receives the class object as `cls` (not an instance).
                   # Allows access to _READERS on the class without constructing an instance.
    def create(cls, ext: str) -> ContextReader:
        """
        Return a ContextReader instance appropriate for the given file extension.

        Parameters
        ----------
        ext : str
            Lowercase file extension including the dot, e.g. ".pdf" or ".txt".

        Returns
        -------
        ContextReader
            A concrete reader that handles files of this type.

        Raises
        ------
        ValueError
            If no reader is registered for the given extension.
        """
        # cls._READERS.get(ext) — dict.get returns None if the key is absent,
        # avoiding a KeyError. In TS: cls._READERS[ext] ?? null
        reader_class = cls._READERS.get(ext)

        if reader_class is None:
            # Build a user-friendly error showing which types ARE supported.
            # ", ".join(...) — join all dict keys with ", " as separator.
            supported = ", ".join(cls._READERS.keys())
            # !r — repr format; wraps the value in quotes in the error message.
            raise ValueError(
                f"Unsupported file type: {ext!r}. Supported types: {supported}"
            )

        # reader_class() — call the class object to construct a new instance.
        # Because reader_class is the class itself (not an instance), calling it
        # with () invokes __init__ and returns the new object.
        return reader_class()


class ContextLoader:
    """
    Public entry point for the document loading layer.

    Responsibility: validate that a file exists, delegate format-specific
    reading to the correct ContextReader (via the factory), and wrap the
    result in the [{filename, content}] dict format consumed by RetrieverBuilder.

    Returning a list (rather than a single dict) lets callers use list.extend()
    to aggregate multiple files into one flat list uniformly.

    Usage:
        loader = ContextLoader()
        docs = loader.load("meeting_notes.pdf")
        # docs == [{"filename": "meeting_notes.pdf", "content": "...full text..."}]
    """

    def load(self, filepath: str) -> list:
        """
        Load a single context file and return its content as a one-element list.

        Parameters
        ----------
        filepath : str
            Path to the context file (.txt, .pdf, or .docx).

        Returns
        -------
        list[dict]
            One-element list: [{"filename": <basename>, "content": <full text>}]
        """
        # os.path.exists(filepath) — returns True if the path exists on disk.
        # In TS: fs.existsSync(filepath)
        if not os.path.exists(filepath):
            print(f"Context file not found: {filepath}")
            sys.exit(1)  # sys.exit(1) — terminate the process with exit code 1 (error)

        _MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB
        if os.path.getsize(filepath) > _MAX_FILE_BYTES:
            print(f"Context file too large (max 50 MB): {filepath}")
            sys.exit(1)

        # os.path.splitext(filepath) — splits "docs/notes.pdf" into ("docs/notes", ".pdf").
        # [1] — index 1 selects the extension part of the tuple.
        # .lower() — normalise extension to lowercase ("PDF" → ".pdf") for dict lookup.
        # In TS: path.extname(filepath).toLowerCase()
        ext: str = os.path.splitext(filepath)[1].lower()

        # os.path.basename(filepath) — strips the directory prefix and returns just
        # the filename: "docs/notes.pdf" → "notes.pdf".
        # In TS: path.basename(filepath)
        filename: str = os.path.basename(filepath)

        try:
            # Ask the factory for the right reader class and instantiate it.
            reader: ContextReader = ContextReaderFactory.create(ext)
            # Delegate the actual file reading to the concrete reader.
            content: str = reader.read(filepath)
        except ValueError as e:
            # ValueError is raised by the factory for unsupported file extensions.
            print(str(e))   # str(e) — convert the exception object to a human-readable string
            sys.exit(1)

        # Return a one-element list so callers can use list.extend() uniformly
        # when aggregating multiple files.
        return [{"filename": filename, "content": content}]
